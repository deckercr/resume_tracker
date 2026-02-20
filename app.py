import io
import os
import sys
import shutil
import tempfile
import webbrowser
from datetime import date, datetime

import webview
from sqlalchemy.orm import Session

# When frozen by PyInstaller, bundled files are in sys._MEIPASS
# The database should live next to the exe, not inside the temp bundle
if getattr(sys, 'frozen', False):
    BUNDLE_DIR = sys._MEIPASS
    APP_DIR = os.path.dirname(sys.executable)
else:
    BUNDLE_DIR = os.path.abspath(os.path.dirname(__file__))
    APP_DIR = BUNDLE_DIR

ASSETS_DIR = os.path.join(BUNDLE_DIR, 'assets')

# Set DB path before importing models so it uses the right location
os.environ['JOB_TRACKER_DB'] = os.path.join(APP_DIR, 'job_tracker.db')

from export import export_to_excel
from models import Application, JobDescription, Resume, engine, init_db


def _find_unicode_fonts():
    """Find a Unicode TTF font (regular + bold) on the system."""
    candidates = [
        # Windows
        ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/arialbd.ttf'),
        ('C:/Windows/Fonts/calibri.ttf', 'C:/Windows/Fonts/calibrib.ttf'),
        # Linux - Noto Sans
        ('/usr/share/fonts/noto/NotoSans-Regular.ttf', '/usr/share/fonts/noto/NotoSans-Bold.ttf'),
        ('/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf', '/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf'),
        # Linux - DejaVu
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
        ('/usr/share/fonts/dejavu/DejaVuSans.ttf', '/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf'),
        # macOS
        ('/System/Library/Fonts/Helvetica.ttc', '/System/Library/Fonts/Helvetica.ttc'),
        ('/Library/Fonts/Arial.ttf', '/Library/Fonts/Arial Bold.ttf'),
    ]
    for regular, bold in candidates:
        if os.path.isfile(regular):
            return regular, bold if os.path.isfile(bold) else regular
    return None, None


class Api:
    # ── Applications ──

    def get_applications(self, status_filter=''):
        with Session(engine) as session:
            query = session.query(Application).order_by(Application.date.desc())
            if status_filter:
                query = query.filter_by(status=status_filter)
            return [self._app_to_dict(a) for a in query.all()]

    def get_application(self, id):
        with Session(engine) as session:
            app = session.get(Application, id)
            if not app:
                return None
            d = self._app_to_dict(app)
            d['has_description'] = app.job_description is not None
            d['resume_name'] = app.resume.name if app.resume else None
            return d

    def create_application(self, data):
        with Session(engine) as session:
            app = Application(
                date=datetime.strptime(data['date'], '%Y-%m-%d').date() if data.get('date') else date.today(),
                status=data.get('status', 'Applied'),
                company=data['company'],
                position=data['position'],
                location=data.get('location', ''),
                url=data.get('url', ''),
                job_id=data.get('job_id', ''),
                salary=data.get('salary', ''),
                contact=data.get('contact', ''),
                notes=data.get('notes', ''),
            )
            session.add(app)
            session.commit()
            return app.id

    def update_application(self, id, data):
        with Session(engine) as session:
            app = session.get(Application, id)
            if not app:
                return False
            if data.get('date'):
                app.date = datetime.strptime(data['date'], '%Y-%m-%d').date()
            app.status = data.get('status', app.status)
            app.company = data.get('company', app.company)
            app.position = data.get('position', app.position)
            app.location = data.get('location', '')
            app.url = data.get('url', '')
            app.job_id = data.get('job_id', '')
            app.salary = data.get('salary', '')
            app.contact = data.get('contact', '')
            app.notes = data.get('notes', '')
            session.commit()
            return True

    def delete_application(self, id):
        with Session(engine) as session:
            app = session.get(Application, id)
            if not app:
                return False
            session.delete(app)
            session.commit()
            return True

    # ── Job Descriptions ──

    def get_job_description(self, application_id):
        with Session(engine) as session:
            jd = session.query(JobDescription).filter_by(application_id=application_id).first()
            if not jd:
                return None
            return {'description': jd.description, 'company': jd.company, 'position': jd.position}

    def save_job_description(self, application_id, description):
        with Session(engine) as session:
            app = session.get(Application, application_id)
            if not app:
                return False
            jd = session.query(JobDescription).filter_by(application_id=application_id).first()
            if jd:
                jd.description = description
                jd.company = app.company
                jd.position = app.position
            else:
                jd = JobDescription(
                    application_id=application_id,
                    company=app.company,
                    position=app.position,
                    description=description,
                )
                session.add(jd)
            session.commit()
            return True

    # ── Resumes ──

    def get_resumes(self):
        with Session(engine) as session:
            resumes = session.query(Resume).order_by(Resume.created_at.desc()).all()
            return [self._resume_to_dict(r) for r in resumes]

    def get_resume(self, id):
        with Session(engine) as session:
            r = session.get(Resume, id)
            if not r:
                return None
            d = self._resume_to_dict(r)
            d['linked_apps'] = [
                {'id': a.id, 'company': a.company, 'position': a.position}
                for a in r.applications
            ]
            return d

    def upload_resume(self, name):
        result = window.create_file_dialog(
            webview.OPEN_DIALOG,
            file_types=('Document Files (*.pdf;*.docx)',),
        )
        if not result:
            return None

        filepath = result[0] if isinstance(result, (list, tuple)) else result
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in ('.pdf', '.docx'):
            return None

        with open(filepath, 'rb') as f:
            file_data = f.read()

        fmt = ext[1:]  # 'pdf' or 'docx'
        with Session(engine) as session:
            resume = Resume(
                name=name or os.path.splitext(os.path.basename(filepath))[0],
                original_format=fmt,
                pdf_data=file_data if fmt == 'pdf' else None,
                docx_data=file_data if fmt == 'docx' else None,
            )
            session.add(resume)
            session.commit()
            return resume.id

    def delete_resume(self, id):
        with Session(engine) as session:
            r = session.get(Resume, id)
            if not r:
                return False
            for app in r.applications:
                app.resume_id = None
            session.delete(r)
            session.commit()
            return True

    def view_resume_file(self, id, fmt):
        with Session(engine) as session:
            r = session.get(Resume, id)
            if not r:
                return False
            data = r.pdf_data if fmt == 'pdf' else r.docx_data
            if not data:
                return False

        suffix = f'.{fmt}'
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(data)
        tmp.close()
        webbrowser.open(tmp.name)
        return True

    def convert_to_pdf(self, id):
        try:
            from docx import Document
            from fpdf import FPDF

            with Session(engine) as session:
                r = session.get(Resume, id)
                if not r or not r.docx_data:
                    return {'ok': False, 'error': 'No DOCX data found'}

                docx_bytes = bytes(r.docx_data)

            doc = Document(io.BytesIO(docx_bytes))
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            font_regular, font_bold = _find_unicode_fonts()
            if font_regular:
                pdf.add_font('UniFont', '', font_regular)
                pdf.add_font('UniFont', 'B', font_bold or font_regular)
                font_name = 'UniFont'
            else:
                font_name = 'Helvetica'

            pdf.set_font(font_name, size=11)

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    pdf.ln(4)
                    continue

                style = ''
                size = 11
                if para.style and para.style.name:
                    sname = para.style.name.lower()
                    if 'heading 1' in sname:
                        style, size = 'B', 16
                    elif 'heading 2' in sname:
                        style, size = 'B', 14
                    elif 'heading 3' in sname:
                        style, size = 'B', 12

                pdf.set_font(font_name, style=style, size=size)
                pdf.multi_cell(0, 6, text)
                pdf.ln(2)

            pdf_bytes = bytes(pdf.output())

            with Session(engine) as session:
                r = session.get(Resume, id)
                r.pdf_data = pdf_bytes
                session.commit()

            return {'ok': True}
        except Exception as e:
            print(f'convert_to_pdf error: {e}')
            return {'ok': False, 'error': str(e)}

    def convert_to_docx(self, id):
        try:
            import fitz
            from docx import Document

            with Session(engine) as session:
                r = session.get(Resume, id)
                if not r or not r.pdf_data:
                    return {'ok': False, 'error': 'No PDF data found'}

                pdf_bytes = bytes(r.pdf_data)

            pdf_doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            doc = Document()

            for page_num in range(len(pdf_doc)):
                if page_num > 0:
                    doc.add_page_break()
                page = pdf_doc[page_num]
                text = page.get_text()
                for line in text.split('\n'):
                    doc.add_paragraph(line)

            pdf_doc.close()

            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()

            with Session(engine) as session:
                r = session.get(Resume, id)
                r.docx_data = docx_bytes
                session.commit()

            return {'ok': True}
        except Exception as e:
            print(f'convert_to_docx error: {e}')
            return {'ok': False, 'error': str(e)}

    def link_resume(self, app_id, resume_id):
        with Session(engine) as session:
            app = session.get(Application, app_id)
            if not app:
                return False
            app.resume_id = resume_id
            session.commit()
            return True

    def unlink_resume(self, app_id):
        with Session(engine) as session:
            app = session.get(Application, app_id)
            if not app:
                return False
            app.resume_id = None
            session.commit()
            return True

    # ── Export ──

    def export_excel(self):
        with Session(engine) as session:
            applications = session.query(Application).order_by(Application.date.desc()).all()
            descriptions = session.query(JobDescription).all()
            tmp_path = export_to_excel(applications, descriptions)

        result = window.create_file_dialog(
            webview.SAVE_DIALOG,
            save_filename='job_applications.xlsx',
            file_types=('Excel Files (*.xlsx)',),
        )
        if result:
            dest = result if isinstance(result, str) else result[0]
            shutil.copy2(tmp_path, dest)
            os.unlink(tmp_path)
            return dest
        os.unlink(tmp_path)
        return None

    def open_url(self, url):
        webbrowser.open(url)

    # ── Helpers ──

    def _app_to_dict(self, app):
        return {
            'id': app.id,
            'date': app.date.strftime('%Y-%m-%d') if app.date else '',
            'status': app.status,
            'company': app.company,
            'position': app.position,
            'location': app.location or '',
            'url': app.url or '',
            'job_id': app.job_id or '',
            'salary': app.salary or '',
            'contact': app.contact or '',
            'notes': app.notes or '',
            'resume_id': app.resume_id,
        }

    def _resume_to_dict(self, r):
        return {
            'id': r.id,
            'name': r.name,
            'original_format': r.original_format,
            'has_pdf': r.pdf_data is not None,
            'has_docx': r.docx_data is not None,
            'created_at': r.created_at.strftime('%Y-%m-%d %H:%M') if r.created_at else '',
            'app_count': len(r.applications),
        }


if __name__ == '__main__':
    init_db()
    api = Api()
    window = webview.create_window(
        'Job Tracker',
        os.path.join(ASSETS_DIR, 'index.html'),
        js_api=api,
        width=1100,
        height=750,
    )
    webview.start(gui='qt')
